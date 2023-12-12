import os
import pathlib
from datetime import date, datetime

import docx
import docx.document
from simplify_docx import simplify
from boltons.iterutils import remap
import pandas as pd

from typing import Any, Dict, Optional, Union

try:
    from plugins.com import word
    IMPORT_PYWIN32COM_OK = True
except ModuleNotFoundError:
    IMPORT_PYWIN32COM_OK = False

class WordRecipes:

    @classmethod
    def read_all(cls, document: Union[str, docx.document.Document]) -> str:
        """
        Extrai o conteúdo de texto de todo o documento.

        :param document: pode ser o caminho do documento ou o objeto do documento
        :returns: string contendo todo o texto do documento
        """
        assert isinstance(document, (str, docx.document.Document)), 'Tipo inválido de documento'

        if isinstance(document, str):
            doc = docx.Document(document)
        else:
            doc = document

        # a biblioteca python-docx permite percorrer parágrafos e tabelas, mas em separado, não dá para saber
        # se um parágrafo veio antes de uma tabela e vice-versa.
        # o pulo do gato está na função simplify, que traz a estrutura do documento em um dicionário organizado.
        document_structure = simplify(doc)

        # o difícil é passar por toda a estrutura, pois é no formato de árvore, por isso usamos remap
        paragraph_texts = []
        previous_kv = []

        def visit(path, key, value):
            if previous_kv:
                previous_k, previous_v = previous_kv[-1]
                if previous_k=='TYPE' and previous_v=='text':
                    paragraph_texts.append(value)
            previous_kv.append((key, value))
            return True

        remap(document_structure, visit=visit)

        # agora junta tudo e retorna
        return '\n'.join(paragraph_texts)
    
    @classmethod
    def table_to_dataframe(cls, document: Union[str, docx.document.Document], table_index: int) -> pd.DataFrame:
        """
        Converte tabela em DataFrame

        :param document: pode ser o caminho do documento ou o objeto do documento
        :param table_index: índice da tabela do documento a ser convertida em DataFrame
        :returns: DataFrame com os dados da tabela
        """
        assert isinstance(document, (str, docx.document.Document)), 'Tipo inválido de documento'

        if isinstance(document, str):
            doc = docx.Document(document)
        else:
            doc = document

        table = doc.tables[table_index]
        cols = [c.text.strip() for c in table.rows[0].cells]
    
        data = []
        # percorre linha a linha
        for row in table.rows[1:]:
            cells = tuple([c.text.strip() for c in row.cells])
            data.append(cells)
        
        return pd.DataFrame(data, columns=cols)
        
    @classmethod
    def fill_template(cls, 
                      filename: str, 
                      save_as: str,
                      values_dict: Dict[str, Any], 
                      *, 
                      placeholder_pattern: Optional[str] = '{{placeholder}}'):
        """
        Preenche um documento substituindo termos por valores de um dicionário.

        :param filename: nome do arquivo de entrada
        :param save_as: nome do arquivo de saída
        :param values_dict: dicionário contendo os termos e os valores a serem substituídos
        :param placeholder_pattern: padrão do placeholder. Note que a palavra "placeholder" deve constar no padrão.
        """
        assert filename.lower().endswith('.docx'), 'Arquivo de entrada precisa ser do tipo DOCX'
        assert save_as.lower().endswith('.docx'), 'Arquivo de saída precisa ser do tipo DOCX'
        assert filename.lower()!=save_as.lower(), 'Os arquivos de entrada e saída não podem ser os mesmos.'
        assert 'placeholder' in placeholder_pattern, 'O padrão de substituição precisa conter a palavra "placeholder"'

        known_converters = {
            float: lambda v: str(v).replace('.', ','),
            date: lambda v: v.strftime(r'%d/%m/%Y'),
            datetime: lambda v: v.strftime(r'%d/%m/%Y'),
        }

        document = docx.Document(filename)

        paragraphs = list(document.paragraphs)
        for t in document.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)

        for p in paragraphs:
            for k, v in values_dict.items():
                placeholder = placeholder_pattern.replace('placeholder', str(k))
                
                v = v or ''
                converted_v = known_converters.get(type(v), str)(v)

                if placeholder in p.text:
                    inline = p.runs
                    
                    # esse tipo de documento é complexo, podemos ter um placeholder distribuído em mais de um run
                    started = False
                    key_index = 0
                    # found_runs é uma lista com (run index, index of match, length of match)
                    found_runs = []
                    found_all = False
                    replace_done = False

                    for i, run in enumerate(inline):

                        # caso 1: encontrou no run inteiro
                        if placeholder in run.text and not started:
                            found_runs.append((i, run.text.find(placeholder), len(placeholder)))
                            text = run.text.replace(placeholder, converted_v)
                            run.text = text
                            replace_done = True
                            found_all = True
                            break

                        if placeholder[key_index] not in run.text and not started:
                            # continua procurando...
                            continue

                        # caso 2: procurar por texto parcial, primeira parte
                        if placeholder[key_index] in run.text and run.text[-1] in placeholder and not started:
                            # verifica sequencia
                            start_index = run.text.find(placeholder[key_index])
                            check_length = len(run.text)
                            for text_index in range(start_index, check_length):
                                if run.text[text_index] != placeholder[key_index]:
                                    # falso positivo
                                    break
                            if key_index == 0:
                                started = True
                            chars_found = check_length - start_index
                            key_index += chars_found
                            found_runs.append((i, start_index, chars_found))
                            if key_index != len(placeholder):
                                continue
                            else:
                                # todos os caracteres foram encontrados
                                found_all = True
                                break

                        # caso 2: procurar por texto parcial, segunda parte
                        if placeholder[key_index] in run.text and started and not found_all:
                            # verifica sequencia
                            chars_found = 0
                            check_length = len(run.text)
                            for text_index in range(0, check_length):
                                if run.text[text_index] == placeholder[key_index]:
                                    key_index += 1
                                    chars_found += 1
                                else:
                                    break
                            # encontrou
                            found_runs.append((i, 0, chars_found))
                            if key_index == len(placeholder):
                                found_all = True
                                break

                    # replace dos runs distribuidos
                    if found_all and not replace_done:
                        for i, item in enumerate(found_runs):
                            index, start, length = [t for t in item]
                            if i == 0:
                                text = inline[index].text.replace(inline[index].text[start:start + length], converted_v)
                                inline[index].text = text
                            else:
                                text = inline[index].text.replace(inline[index].text[start:start + length], '')
                                inline[index].text = text

        if os.path.exists(save_as):
            os.remove(save_as)

        document.save(save_as)

    @classmethod
    def to_pdf(cls, doc_filename: str, pdf_filename: str):
        """
        Converte arquivo DOCX em PDF. Necessário ter o Word instalado.

        :param doc_filename: nome do arquivo DOCX de entrada
        :param pdf_filename: nome do arquivo PDF de saída
        """
        assert IMPORT_PYWIN32COM_OK, 'Os requisitos para essa funcionalidade não foram instalados/ configurados.'
        
        with word.application(kill_after=True) as wd:
            doc_filename = str(pathlib.Path(doc_filename).resolve())
            pdf_filename = str(pathlib.Path(pdf_filename).resolve())

            doc = wd.Documents.Open(doc_filename)

            if os.path.exists(pdf_filename):
                os.remove(pdf_filename)

            doc.SaveAs(pdf_filename, FileFormat=word.constants.wdFormatPDF)
            doc.Close()
        return